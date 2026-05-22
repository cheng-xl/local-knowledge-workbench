import os
import hashlib
from typing import List, Optional, Dict
from rank_bm25 import BM25Okapi
from vector_store import VectorStoreManager
from shared import Document
from config import settings
from loguru import logger


# ── Document loaders ────────────────────────────────────────

def _load_pdf(file_path: str) -> List[Document]:
    from pypdf import PdfReader
    reader = PdfReader(file_path)
    docs = []
    fname = os.path.basename(file_path)
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text and text.strip():
            docs.append(Document(page_content=text.strip(),
                        metadata={"source_file": fname, "file_type": "pdf",
                                  "page": i + 1}))
    return docs


def _load_docx(file_path: str) -> List[Document]:
    from docx import Document as DocxDocument
    doc = DocxDocument(file_path)
    text = "\n".join(p.text for p in doc.paragraphs if p.text)
    fname = os.path.basename(file_path)
    return [Document(page_content=text,
                     metadata={"source_file": fname, "file_type": "docx"})]


def _load_text(file_path: str) -> List[Document]:
    fname = os.path.basename(file_path)
    ext = file_path.rsplit(".", 1)[-1].lower()
    with open(file_path, "r", encoding="utf-8") as f:
        text = f.read()
    return [Document(page_content=text,
                     metadata={"source_file": fname, "file_type": ext})]


_LOADERS: Dict[str, callable] = {
    "pdf": _load_pdf,
    "docx": _load_docx,
    "md": _load_text,
    "txt": _load_text,
}


# ── Text splitter ───────────────────────────────────────────

class RecursiveCharacterTextSplitter:
    def __init__(self, chunk_size: int = 512, chunk_overlap: int = 128,
                 separators: list = None):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.separators = separators or ["\n\n", "\n", "。", ".", " ", ""]

    def _split_by_sep(self, text: str, sep: str) -> List[str]:
        """Split text by a separator, merging pieces up to chunk_size."""
        chunks = []
        current = ""
        for part in text.split(sep):
            candidate = (current + sep + part).strip() if current else part
            if len(candidate) <= self.chunk_size:
                current = candidate
            else:
                if current:
                    chunks.append(current)
                current = part
        if current:
            chunks.append(current)
        return chunks

    def split_text(self, text: str) -> List[str]:
        if len(text) <= self.chunk_size:
            return [text]
        chunks = []
        current = ""
        for para in text.split(self.separators[0]):
            candidate = (current + self.separators[0] + para).strip() if current else para
            if len(candidate) <= self.chunk_size:
                current = candidate
            else:
                if current:
                    chunks.append(current)
                if len(para) > self.chunk_size:
                    chunks.extend(self._split_long(para))
                    current = ""
                else:
                    current = para
        if current:
            chunks.append(current)
        return chunks

    def _split_long(self, text: str) -> List[str]:
        for sep in self.separators[1:]:
            if sep and sep in text:
                return self._split_by_sep(text, sep)
        step = self.chunk_size - self.chunk_overlap
        return [text[i:i + self.chunk_size] for i in range(0, len(text), step)]

    def split_documents(self, documents: List[Document]) -> List[Document]:
        chunks = []
        for doc in documents:
            texts = self.split_text(doc.page_content)
            for text in texts:
                new_doc = Document(
                    page_content=text,
                    metadata=dict(doc.metadata),
                )
                chunks.append(new_doc)
        return chunks


# ── RAG Pipeline ────────────────────────────────────────────

class RAGPipeline:
    def __init__(self, persist_dir: Optional[str] = None,
                 vector_store: "VectorStoreManager" = None):
        self.vs = vector_store or VectorStoreManager(persist_dir)
        self._doc_texts: List[str] = []
        self._bm25: Optional[BM25Okapi] = None
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
        )

    def load_document(self, file_path: str) -> List[Document]:
        ext = file_path.rsplit(".", 1)[-1].lower()
        loader = _LOADERS.get(ext, _load_text)
        logger.info(f"Loading {file_path} ({ext})")
        return loader(file_path)

    def chunk_documents(self, docs: List[Document],
                        chunk_size: int = None,
                        overlap: int = None) -> List[Document]:
        cs = chunk_size or settings.chunk_size
        ov = overlap or settings.chunk_overlap
        if cs != self.splitter.chunk_size or ov != self.splitter.chunk_overlap:
            self.splitter = RecursiveCharacterTextSplitter(
                chunk_size=cs, chunk_overlap=ov)
        chunks = self.splitter.split_documents(docs)
        for i, c in enumerate(chunks):
            c.metadata["chunk_index"] = i
            c.metadata["chunk_hash"] = hashlib.md5(
                c.page_content.encode()).hexdigest()[:8]
        logger.info(f"Split {len(docs)} docs into {len(chunks)} chunks")
        return chunks

    def add_to_store(self, chunks: List[Document]) -> List[str]:
        ids = self.vs.add_documents(chunks)
        self._doc_texts.extend(c.page_content for c in chunks)
        self._bm25 = None
        return ids

    @staticmethod
    def _tokenize(text: str) -> List[str]:
        try:
            import jieba
            return list(jieba.cut(text))
        except ImportError:
            return list(text)

    @property
    def bm25(self) -> Optional[BM25Okapi]:
        if self._bm25 is None and self._doc_texts:
            logger.info(f"Building BM25 index on {len(self._doc_texts)} chunks")
            corpus = [self._tokenize(t) for t in self._doc_texts]
            self._bm25 = BM25Okapi(corpus)
        return self._bm25

    def hybrid_search(self, query: str, top_k: int = None,
                      use_rerank: bool = True) -> List[Document]:
        k = top_k or settings.retrieval_top_k

        vec_results = self.vs.search(query, top_k=k * 2)
        logger.debug(f"Vector search returned {len(vec_results)} results")

        bm25_results = []
        if self.bm25 is not None:
            scores = self.bm25.get_scores(self._tokenize(query))
            top_idx = sorted(range(len(scores)),
                             key=lambda i: scores[i], reverse=True)[:k * 2]
            bm25_results = [
                Document(page_content=self._doc_texts[i],
                         metadata={"bm25_score": float(scores[i]),
                                   "retrieval_method": "bm25"})
                for i in top_idx if scores[i] > 0
            ]
            logger.debug(f"BM25 returned {len(bm25_results)} results")

        merged = self._rrf(vec_results, bm25_results)
        if use_rerank and len(merged) > k:
            merged = self._rerank(query, merged)
        return merged[:k]

    @staticmethod
    def _rrf(vec: List[Document], bm25: List[Document],
             k: int = 60) -> List[Document]:
        scores: dict = {}
        for rank, doc in enumerate(vec):
            key = doc.page_content[:200]
            scores[key] = (doc, 1.0 / (k + rank + 1))
        for rank, doc in enumerate(bm25):
            key = doc.page_content[:200]
            rrf = 1.0 / (k + rank + 1)
            if key in scores:
                scores[key] = (scores[key][0], scores[key][1] + rrf)
            else:
                scores[key] = (doc, rrf)
        sorted_items = sorted(scores.values(), key=lambda x: x[1], reverse=True)
        return [item[0] for item in sorted_items]

    @staticmethod
    def _rerank(query: str, docs: List[Document]) -> List[Document]:
        try:
            from FlagEmbedding import FlagReranker
            reranker = FlagReranker("BAAI/bge-reranker-base")
            pairs = [[query, d.page_content] for d in docs]
            scores = reranker.compute_score(pairs, normalize=True)
            scored = sorted(zip(docs, scores), key=lambda x: x[1], reverse=True)
            logger.debug(f"Reranking: top score={scored[0][1]:.3f}")
            return [d for d, _ in scored]
        except ImportError:
            logger.warning("FlagEmbedding not installed, skipping rerank")
            return docs

    def ingest_file(self, file_path: str,
                    chunk_size: int = None, overlap: int = None) -> int:
        docs = self.load_document(file_path)
        chunks = self.chunk_documents(docs, chunk_size=chunk_size, overlap=overlap)
        self.add_to_store(chunks)
        return len(chunks)
